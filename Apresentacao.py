# /home/neirivon/SINAPSE2.0/sinapsebr_rubrica/Apresentacao.py
# --------------------------------------------------------------------------------------
# NOME DO SCRIPT: Apresentacao.py
# DESCRIÇÃO: Página inicial (Home) do protótipo SINAPSE-BR IA.
#            Apresenta o resumo estruturado do TCC, equipe, fluxo metodológico e opções de exportação.
#
# FUNCIONALIDADES:
#   1. Resumo Expandido: Problema, Objetivos, Metodologia e Resultados.
#   2. Perfis: Cards visuais com efeito Mouseover e caminhos corrigidos.
#   3. Diagrama DBR: Visualização do fluxo da pesquisa (Graphviz).
#   4. Exportação: Geração dinâmica de DOCX e PDF (Corrigida para evitar "?" no texto).
#
# AUTOR: Neirivon Elias Cardoso
# PROJETO: Rubrica SINAPSE-BR IA — Sistema Integrado Neuropsicopedagógico
# TCC: Pós-Graduação em Docência para a EPT (IFTM)
# DATA: 07/02/2026 (Update: Correção de Encoding no PDF)
# --------------------------------------------------------------------------------------

import os
import io
import re
import base64
import graphviz 
from pathlib import Path
from io import BytesIO
from PIL import Image, ImageDraw
import streamlit as st

# --- IMPORTS PARA EXPORTAÇÃO ---
try:
    from docx import Document
    from fpdf import FPDF
except ImportError:
    st.error("Bibliotecas 'python-docx' ou 'fpdf' não instaladas. As funções de exportação não funcionarão.")

# ---------------------------------
# Config da página
# ---------------------------------
st.set_page_config(
    page_title="SINAPSE-BR IA — Apresentação TCC",
    page_icon="🧠",
    layout="wide",
)

# ---------------------------------
# Localizador robusto (assets/imagens)
# ---------------------------------
THIS = Path(__file__).resolve()

def find_project_root(start: Path) -> Path:
    p = start
    for _ in range(6): 
        if (p / "assets").exists():
            return p
        if p.parent == p:
            break
        p = p.parent
    return Path.cwd()

PROJECT_ROOT = find_project_root(THIS.parent)
IMAGENS_DIR = PROJECT_ROOT / "assets" / "imagens"

# Caminhos específicos
NEIRIVON_IMG      = IMAGENS_DIR / "neirivon.png"
ORIENTADORA_IMG   = IMAGENS_DIR / "Orientadora.png"
LOGO_IFTM         = IMAGENS_DIR / "IFTM_360.png"
LOGO_SINAPSE      = IMAGENS_DIR / "sinapse.png"

# ---------------------------------
# Utilitários de imagem / HTML
# ---------------------------------

def img_circular_b64(path: Path) -> str:
    try:
        if not path.exists(): return ""
        img = Image.open(path).convert("RGBA")
        mask = Image.new('L', img.size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, img.size[0], img.size[1]), fill=255)
        output_buffer = BytesIO()
        img.putalpha(mask)
        img.save(output_buffer, format="PNG")
        output_buffer.seek(0)
        return f"data:image/png;base64,{base64.b64encode(output_buffer.read()).decode()}"
    except Exception: return ""

def tag_html_profile_content(base64_img: str, name: str, caption: str):
    return f"""
    <div class="profile-container">
        <img class="profile-image" src="{base64_img}" alt="{name}">
        <div class="profile-details">
            <h3>{name}</h3>
            <p>{caption}</p>
        </div>
    </div>
    """

def safe_image(path: Path, *, width: int | None = None, caption: str | None = None):
    try:
        if path.exists(): st.image(str(path), width=width, caption=caption)
    except Exception as e: st.error(f"Erro imagem: {e}")

# ---------------------------------
# DIAGRAMA METODOLÓGICO (DBR)
# ---------------------------------
def render_dbr_diagram():
    """Gera o diagrama de fluxo da pesquisa usando Graphviz"""
    graph = graphviz.Digraph()
    graph.attr(rankdir='LR', bgcolor='transparent')
    graph.attr('node', shape='box', style='rounded,filled', fontname='Helvetica')
    
    # Definição dos Nós
    graph.node('P', 'PROBLEMA\n(Exclusão na EPT)', fillcolor='#ffcccc', color='#ff0000')
    graph.node('T', 'DESIGN TEÓRICO\n(Neuro + DUA + Território)', fillcolor='#ccffcc', color='#009900')
    graph.node('D', 'PROTOTIPAGEM\n(Rubrica SINAPSE + IA)', fillcolor='#ccccff', color='#0000ff')
    graph.node('V', 'VALICAÇÃO\n(Meta-Rubrica)', fillcolor='#ffffcc', color='#ffcc00')
    
    # Definição das Arestas
    graph.edge('P', 'T', label=' Análise')
    graph.edge('T', 'D', label=' Desenvolvimento')
    graph.edge('D', 'V', label=' Avaliação')
    graph.edge('V', 'T', label=' Refinamento (Ciclo DBR)', style='dashed', color='grey')
    
    st.graphviz_chart(graph, use_container_width=True)

# ---------------------------------
# FUNÇÕES DE GERAÇÃO DE ARQUIVOS
# ---------------------------------

def get_apresentacao_content():
    content = []
    content.append("# 🧠 SINAPSE-BR IA")
    content.append("## Sistema Integrado Neuropsicopedagógico de Avaliação e Práticas da Sinergia Educacional Para a EPT")
    content.append("Operando sob o paradigma da Inteligência Artificial Centrada no Humano, com supervisão técnica e pedagógica contínua, garantindo que a tecnologia atue como um amplificador da expertise docente, nunca como substituta do julgamento clínico-pedagógico.")
    content.append("---")
    content.append("## 🧑‍🎓 Autoria e Orientação")
    content.append("**Orientando:** Neirivon Elias Cardoso (Pós-graduando em Docência para a EPT)")
    content.append("**Orientadora:** Prof.ª Dra. Thays Martins Vital da Silva (Professora Orientadora)")
    content.append("**Instituição:** IFTM - Campus Avançado Uberaba Parque Tecnológico")
    content.append("**Curso:** Pós-Graduação Lato Sensu em Docência para a EPT")
    content.append("---")
    content.append("## 📚 Núcleo da Pesquisa")
    content.append("### TEMA")
    content.append("> Desenvolvimento de rubrica educacional ampliada para avaliação formativa na EPT, integrando Neuropsicopedagogia, DUA e Inteligência Artificial.")
    content.append("### PROBLEMA")
    content.append("`Como integrar princípios da Neuropsicopedagogia, das Taxonomias Cognitivas (Bloom/SOLO), do DUA e da equidade socioterritorial em uma rubrica formativa aplicável à EPT?`")
    content.append("O problema reside na invisibilidade das barreiras neurocognitivas e sociais, buscando superar a exclusão denunciada por Ciavatta e Ramos.")
    content.append("### DELIMITAÇÃO")
    content.append("Construção teórico-propositiva da **Rubrica SINAPSE-BR IA** para a Rede Federal, com recorte territorial e analítico no **Triângulo Mineiro e Alto Paranaíba (TMAP)**, utilizando dados do SISTEC, Censo Escolar e o **resgate da memória institucional**.")
    content.append("---")
    content.append("## 1. Introdução e Justificativa ✍️")
    content.append("A avaliação na EPT enfrenta o desafio de superar a simples verificação de tarefas técnicas. A proposta combate a visão tecnicista da avaliação denunciada por Ciavatta e Ramos, propondo um modelo que considera como o cérebro aprende (Neurociência) e onde o aluno está (Território).")
    content.append("### 🎯 Objetivo Geral")
    content.append("Analisar e propor uma rubrica educacional ampliada — SINAPSE-BR IA — fundamentada em referenciais de Neuropsicopedagogia, DUA e equidade territorial.")
    content.append("### 🎯 Objetivos Específicos")
    content.append("* **1.** Analisar referenciais teóricos: Neuropsicopedagogia, Taxonomias e modelos de avaliação da EPT.\n* **2.** Comparar estruturas de rubricas nacionais e internacionais.\n* **3.** Propor a estrutura da Rubrica SINAPSE-BR IA.")
    content.append("---")
    content.append("## 2. Fundamentação Teórica 🏗️")
    content.append("### O que é EPT?")
    content.append("Segundo Ciavatta, Ramos e Frigotto (2005), a EPT transcende o mero treinamento técnico; deve ser compreendida como uma formação humana integral que articula Trabalho, Ciência, Cultura e Tecnologia, assumindo o trabalho como princípio educativo.")
    content.append("### O que é uma Rubrica?")
    content.append("Segundo Susan Brookhart (2013), uma rubrica não é uma simples checklist, mas um 'conjunto coerente de critérios para o trabalho dos alunos que inclui descrições dos níveis de qualidade'. Foca no resultado de aprendizagem e feedback descritivo.")
    content.append("### O que é a IA Intermediada por Humanos?")
    content.append("Paradigma Human-in-the-loop: a IA atua como 'Intelecto Objetivado' (Saviani), organizando dados complexos (SISTEC/Território), mas a decisão pedagógica permanece sob controle soberano do docente.")
    content.append("1. **EPT e Trabalho:** Saviani, Ramos, Frigotto.\n2. **Neuropsicopedagogia:** Cosenza & Guerra, Piaget.\n3. **Inclusão e DUA:** Desenho Universal para a Aprendizagem.\n4. **Territorialização:** Dados do SISTEC e Censo Escolar.")
    content.append("---")
    content.append("## 3. Metodologia (Design-Based Research) 🧪")
    content.append("A Design-Based Research (DBR) permite criar soluções práticas fundamentadas em teoria e testá-las ciclicamente no ambiente real do IFTM. Natureza: Pesquisa Aplicada, abordagem Qualitativa e cunho Teórico-Propositiva.")
    content.append("### Procedimentos:")
    content.append("* **Análise Preliminar:** Diagnóstico da exclusão avaliativa e levantamento bibliográfico.")
    content.append("* **Design Teórico:** Engenharia Pedagógica (Matriz Tridimensional).")
    content.append("* **Prototipagem:** Desenvolvimento do artefato em Python/Streamlit.")
    content.append("* **Avaliação:** Validação por Meta-Rubrica.")
    content.append("---")
    content.append("## 4. Produto Educacional 🖥️")
    content.append("- **Rubrica Interpretativa:** Matriz tridimensional de avaliação.\n- **Painel Territorial:** Uso de dados do INEP/SISTEC para revelar desigualdades no Triângulo Mineiro.\n- **Relatórios Automatizados:** Feedbacks humanos potencializados por IA.")
    content.append("---")
    content.append("## 5. Considerações Finais ✅")
    content.append("A SINAPSE-BR IA não é apenas uma ferramenta de notas, mas um sistema de apoio à decisão que integra **dados** e **pedagogia**, gerando princípios de design para uma EPT mais inclusiva.")
    return "\n\n".join(content)

def clean_text_for_pdf(text):
    """Remove emojis e caracteres que causam erro '?' no PDF"""
    # Remove emojis usando regex
    text = re.sub(r'[^\x00-\x7F]+', ' ', text)
    # Remove aspas curvas e outros caracteres latinos especiais que falham no latin-1
    text = text.replace('‘', "'").replace('’', "'").replace('“', '"').replace('”', '"')
    return text

def generate_docx(content_markdown: str) -> BytesIO:
    document = Document()
    document.add_heading("SINAPSE-BR IA - Apresentação TCC", 0)
    for line in content_markdown.split('\n\n'):
        line = line.strip()
        if not line: continue
        if line.startswith('## '): document.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '): document.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('>'): 
            p = document.add_paragraph(line.replace('>', '').strip())
            p.style = 'Intense Quote'
        elif line.startswith('`'): document.add_paragraph(line.replace('`', '').strip())
        elif '*' in line:
             for item in line.split('\n'):
                if item.strip(): document.add_paragraph(item.replace('*', '').strip(), style='List Bullet')
        else: document.add_paragraph(line)
    doc_buffer = BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def generate_pdf(content_markdown: str) -> BytesIO:
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 12)
            self.cell(0, 10, 'SINAPSE-BR IA - Apresentacao TCC', 0, 1, 'C')
            self.ln(5)
        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Pagina {self.page_no()}/{{nb}}', 0, 0, 'C')
    
    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    for line in content_markdown.split('\n\n'):
        line = line.strip()
        if not line: pdf.ln(2); continue
        
        # Limpamos a linha de emojis para evitar o erro do "?"
        line_clean = clean_text_for_pdf(line)
        
        if line.startswith('##'):
            pdf.set_font('Arial', 'B', 14 if line.count('#')==1 else 12)
            pdf.multi_cell(0, 8, line_clean.replace('#', '').strip())
            pdf.ln(2)
        else:
            pdf.set_font('Arial', '', 10)
            pdf.multi_cell(0, 6, line_clean.replace('*', '').strip())
            pdf.ln(1)
            
    return BytesIO(pdf.output(dest='S'))

# ---------------------------------
# SIDEBAR
# ---------------------------------
with st.sidebar:
    safe_image(LOGO_IFTM, width=130) 
    st.markdown("---")
    st.markdown("### 📑 Navegação TCC")
    st.caption("Apresentação do Trabalho de Conclusão de Curso.")
    st.markdown("---")
    safe_image(LOGO_SINAPSE, width=200)
    st.caption("SINAPSE-BR • Sistema Integrado Neuropsicopedagógico")
    st.markdown("---")
    
# ---------------------------------
# PÁGINA PRINCIPAL
# ---------------------------------

st.title("🧠 SINAPSE-BR IA")
st.subheader("Sistema Integrado Neuropsicopedagógico de Avaliação e Práticas da Sinergia Educacional Para a EPT")

# Texto de alto nível sobre supervisão humana solicitado pelo autor
st.markdown("""
*Operando sob o paradigma da **Inteligência Artificial Centrada no Humano**, com supervisão técnica e pedagógica contínua, 
garantindo que a tecnologia atue como um amplificador da expertise docente, nunca como substituta do julgamento clínico-pedagógico.*
""")

st.markdown("**Curso:** Pós-Graduação em Docência para a EPT — IFTM")

col_side, col_main = st.columns([0.25, 0.75])

with col_main:
    # CSS: Efeito Mouseover e estilos gerais
    st.markdown("""
    <style>
        .profile-container {
            display: flex; align-items: center; padding: 15px;
            background-color: #f8f9fa; border-radius: 10px;
            margin-bottom: 15px; border: 1px solid #e0e0e0;
            transition: all 0.3s ease;
        }
        .profile-container:hover {
            box-shadow: 0 4px 12px rgba(0,0,0,0.1); border-color: #b0bec5;
        }
        .profile-image {
            width: 120px; height: 120px; border-radius: 50%;
            object-fit: cover; margin-right: 20px;
            border: 3px solid #4CAF50;
            transition: transform 0.3s ease, box-shadow 0.3s ease;
        }
        .profile-image:hover {
            transform: scale(1.1); box-shadow: 0 0 15px rgba(76, 175, 80, 0.6); cursor: pointer;
        }
        .profile-details h3 { margin: 0 0 5px 0; color: #2c3e50; font-size: 1.3em;}
        .profile-details p { margin: 0; color: #7f8c8d; font-size: 0.95em;}
    </style>
    """, unsafe_allow_html=True)

    st.markdown("### 🧑‍🎓 Autoria e Orientação")
    if NEIRIVON_IMG.exists():
        st.markdown(tag_html_profile_content(img_circular_b64(NEIRIVON_IMG), "Neirivon Elias Cardoso", "Pós-graduando em Docência para a EPT"), unsafe_allow_html=True)
    if ORIENTADORA_IMG.exists():
        st.markdown(tag_html_profile_content(img_circular_b64(ORIENTADORA_IMG), "Dra. Thays Martins Vital da Silva", "Orientadora do TCC"), unsafe_allow_html=True)

    st.divider()

    # --- NÚCLEO DA PESQUISA (Visual Aprimorado) ---
    with st.expander("📚 Núcleo da Pesquisa (Problema e Delimitação)", expanded=True):
        
        st.markdown("##### 🎯 Tema")
        st.info("**Desenvolvimento de rubrica educacional ampliada para EPT com base em Neuropsicopedagogia e IA.**")
        
        st.markdown("##### ❓ Problema de Pesquisa")
        st.warning(
            "**Como integrar princípios da Neuropsicopedagogia, das Taxonomias Cognitivas (Bloom/SOLO), "
            "do DUA e da equidade socioterritorial em uma rubrica formativa aplicável à EPT?**",
            icon="🤔"
        )
        with st.expander("📖 Por que isso é um problema?"):
            st.write("""
            A avaliação na EPT muitas vezes foca apenas na competência técnica imediata. 
            O problema reside na invisibilidade das barreiras neurocognitivas e sociais. 
            A pesquisa questiona como criar um instrumento (rubrica) que seja justo e integral, superando a exclusão denunciada por Ciavatta e Ramos.
            """)
        
        st.markdown("##### 🔭 Delimitação do Estudo")
        st.markdown("Foco na **Rede Federal (IFTM)**, recorte territorial **TMAP**, utilizando:")
        st.markdown("""
        * 📊 **Dados quantitativos:** do SISTEC e Censo Escolar;
        * 🧠 **Resgate da memória institucional** e dados qualitativos do território.
        """)

    # --- ABAS DE DETALHAMENTO (CONTEÚDO ENRIQUECIDO COM O TCC) ---
    tab1, tab2, tab3, tab4 = st.tabs(["Justificativa", "Teoria", "Metodologia", "Produto"])
    
    with tab1:
        st.markdown("### 🚀 Por que SINAPSE-BR?")
        st.write("A avaliação na EPT enfrenta o desafio de superar a dualidade estrutural. A proposta cria um **Ecossistema Avaliativo** que integra a **Neurociência** (como se aprende) e a **Geofilosofia** (onde se vive).")
        with st.expander("📖 Expandir Justificativa"):
            st.write("Buscamos promover a **Permeabilidade Seletiva** entre a escola e o território, combatendo o adestramento técnico e instrumentando o docente para uma avaliação que reconhece o sujeito integral.")
        st.success("**Objetivo:** Instrumentar o docente para uma avaliação que reconhece o sujeito integral.")

    with tab2:
        st.markdown("### 🏗️ Fundamentação Teórica")

        # Inclusão da Definição de EPT conforme solicitado
        with st.expander("📖 1. O que é EPT? (Ciavatta/Ramos/Frigotto)", expanded=False):
            st.write("""
            A Educação Profissional e Tecnológica (EPT) transcende o mero treinamento técnico para o mercado. 
            Segundo **Ciavatta, Ramos e Frigotto (2005)**, ela deve ser compreendida como uma formação humana integral que articula Trabalho, Ciência, Cultura e Tecnologia. 
            Nesta perspectiva, o trabalho é assumido como princípio educativo (ontológico) e não apenas econômico, visando à superação da dualidade histórica entre fazer manual e saber intelectual, rumo à politecnia e à emancipação dos sujeitos.
            """)
        
        with st.expander("📖 2. O que é uma Rubrica? (Brookhart)", expanded=False):
            st.write("""
            Segundo **Susan Brookhart (2013)**, uma rubrica não é uma simples checklist. Ela é um 'conjunto coerente de critérios para o trabalho dos alunos que inclui descrições dos níveis de qualidade'. 
            Foca no resultado de aprendizagem e feedback descritivo para que o estudante entenda sua zona de desenvolvimento proximal.
            """)

        with st.expander("📖 3. IA Intermediada por Humanos (Russell/Nicolelis)", expanded=False):
            st.markdown("""
            A Rubrica SINAPSE-BR IA opera sob o paradigma <span title="TRADUÇÃO: Humano no Ciclo. &#10;EXPLICAÇÃO: Modelo onde a IA requer supervisão humana para treinar, validar e refinar decisões, garantindo que o controle pedagógico e ético permaneça com o docente." style="border-bottom: 1px dotted; cursor: help; font-weight: bold;">Human-in-the-loop</span>.
            A IA atua como 'Intelecto Objetivado' (Saviani/Marx), organizando 
            grandes volumes de dados (SISTEC/Território), mas a decisão pedagógica 
            permanece sob controle humano, garantindo a emancipação e não o 
            adestramento algorítmico.
            """, unsafe_allow_html=True)

        c1, c2, c3 = st.columns(3)
        with c1: 
            st.info("**EPT e Trabalho**\n\nSaviani, Frigotto e Ciavatta.\n*Politecnia e Trabalho como Princípio Educativo.*")
        with c2: 
            st.info("**Neuropsicopedagogia**\n\nCosenza, Guerra e Taxonomias (Bloom/SOLO).\n*Funções Executivas e Níveis Cognitivos.*")
        with c3:
            st.info("**Geofilosofia**\n\nMilton Santos e Paulo Irineu.\n*Território Usado e Identidade Local.*")
    
    with tab3:
        st.markdown("### 🧪 Metodologia (Design-Based Research)")
        st.markdown("""
        Esta pesquisa configura-se como um **Estudo de Desenvolvimento** (*Plomp; Nieveen, 2013*), que visa criar uma intervenção prática (a Rubrica).
        """)
        
        with st.expander("🔄 Ver Fluxo Metodológico (Diagrama)"):
            render_dbr_diagram()
            st.caption("Figura: Ciclo metodológico da Pesquisa de Design Educacional aplicado ao SINAPSE.")
        
        with st.expander("📖 Por que DBR?"):
            st.write("""
            A DBR foi escolhida porque permite criar uma solução (a Rubrica SINAPSE) e testá-la ciclicamente no ambiente real do IFTM. 
            Garante que o produto final seja útil para o cotidiano escolar através do refinamento constante.
            """)
        
        st.markdown("""
        **Fases do Ciclo:**
        1. **Análise Preliminar** | 2. **Design Teórico** | 3. **Prototipagem** | 4. **Validação**
        """)
    
    with tab4:
        st.markdown("### 🖥️ O Produto Educacional")
        st.write("O software materializa a **Matriz Tridimensional de Avaliação**:")
        st.markdown("""
        * 🧠 **Dimensão Cognitiva:** Avalia o processamento mental e a metacognição.
        * 🛠️ **Dimensão da Práxis:** Avalia o domínio técnico e tecnológico.
        * 🌍 **Dimensão Territorial:** Avalia a conexão com a identidade local (Pertencimento).
        """)
        with st.expander("📖 Detalhes Técnicos e Territoriais"):
            st.write("O painel utiliza dados do INEP/SISTEC para revelar desigualdades regionais no Triângulo Mineiro, auxiliando na criação de rubricas que respeitem a realidade local do estudante.")
        st.caption("Baseado em Llama 3 e Python/Streamlit.")

    st.divider()
    st.subheader("⬇️ Obter cópia da Apresentação")
    md_content = get_apresentacao_content()
    c1, c2, c3 = st.columns([1, 1, 2])
    with c1: st.download_button("📄 Salva no formato DOCX", data=generate_docx(md_content), file_name="SINAPSE_TCC.docx")
    with c2: 
        try: st.download_button("📑 Salvar no formato PDF", data=generate_pdf(md_content), file_name="SINAPSE_TCC.pdf")
        except: st.error("Erro PDF")
    
    st.caption(f"Sistema rodando a partir de: `{PROJECT_ROOT}`")
