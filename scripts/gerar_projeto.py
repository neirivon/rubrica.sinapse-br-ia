from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_docx_mestrado():
    doc = Document()

    # --- CONFIGURAÇÃO DE MARGENS (Edital: Sup/Esq 3cm, Inf/Dir 2cm) ---
    for section in doc.sections:
        section.top_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.bottom_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # --- ESTILO PADRÃO (Times New Roman, 12pt, Espaçamento 1,5) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)

    # --- CAPA (ANÓNIMA) ---
    capa = doc.add_paragraph('\n\n\n')
    capa.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_capa = capa.add_run("ECOSSISTEMA DE INTELIGÊNCIA ARTIFICIAL PARA AVALIAÇÃO TRIDIMENSIONAL NA EDUCAÇÃO PROFISSIONAL: UMA ABORDAGEM GEOFILOSÓFICA E MULTIAGENTE")
    titulo_capa.bold = True
    
    doc.add_paragraph('\n\n')
    detalhes = doc.add_paragraph()
    detalhes.alignment = WD_ALIGN_PARAGRAPH.CENTER
    detalhes.add_run("NÚMERO DA LINHA DE PESQUISA: [Inserir aqui]\n")
    detalhes.add_run("NOME DA LINHA DE PESQUISA: [Inserir aqui]\n")
    detalhes.add_run("NÚMERO DE INSCRIÇÃO: [Inserir aqui]")

    doc.add_page_break()

    # --- CONTEÚDO ---
    
    add_heading("1 INTRODUÇÃO")
    doc.add_paragraph(
        "A avaliação na Educação Profissional e Tecnológica (EPT) enfrenta historicamente o desafio de superar modelos lineares e puramente quantitativos, "
        "que frequentemente resultam em um fenômeno de 'achatamento' da subjetividade do educando. No contexto da dualidade estrutural entre o saber "
        "técnico e o saber formativo, emerge a necessidade de mecanismos que considerem a complexidade do sujeito. Esta proposta introduz a concepção "
        "de um ecossistema de Inteligência Artificial (IA) fundamentado na Sinergia Educacional, atuando não como substituto da docência, mas como "
        "mediador neuropsicopedagógico capaz de interpretar rubricas multidimensionais."
    )

    add_heading("2 JUSTIFICATIVA")
    doc.add_paragraph(
        "A presente pesquisa justifica-se pela urgência de inovação social e tecnológica no processo avaliativo da rede federal de ensino. No âmbito social, "
        "busca-se mitigar a redução do desempenho discente a métricas frias, valorizando a trajetória integral do aluno. Tecnicamente, a utilização de "
        "fluxos agênticos (Agentic Workflows) propicia uma transparência inédita em sistemas automatizados, permitindo a auditabilidade do processo. "
        "Sob a ótica da territorialidade, a proposta ancora-se na Geofilosofia para situar o conhecimento no 'espaço usado', elemento essencial para a "
        "identidade dos Institutos Federais em sua relação com o desenvolvimento regional."
    )

    add_heading("3 PROBLEMÁTICA")
    doc.add_paragraph(
        "Diante do cenário de opacidade de muitos algoritmos educacionais contemporâneos, formula-se o seguinte problema: como superar a dualidade "
        "estrutural da EPT e a invisibilidade do território do aluno por meio de um artefato de inteligência artificial que seja transparente, "
        "rastreável e fundamentado em rubricas multidimensionais?"
    )

    add_heading("4 OBJETIVOS")
    p_geral = doc.add_paragraph()
    p_geral.add_run("4.1 OBJETIVO GERAL").bold = True
    doc.add_paragraph("Desenvolver e validar um orquestrador multiagente de inteligência artificial (denominado operacionalmente como ROTOR) focado na regulação de rubricas de avaliação tridimensional no âmbito da Educação Profissional e Tecnológica.")
    
    p_esp = doc.add_paragraph()
    p_esp.add_run("4.2 OBJETIVOS ESPECÍFICOS").bold = True
    doc.add_paragraph("• Mapear os fundamentos da Geofilosofia da Mente e da Neuropsicopedagogia aplicados aos processos avaliativos contemporâneos;", style='List Bullet')
    doc.add_paragraph("• Implementar uma arquitetura de Geração Aumentada de Recuperação (Retrieval-Augmented Generation - RAG) com organização hierárquica e temporal;", style='List Bullet')
    doc.add_paragraph("• Avaliar a eficácia de uma interface de governança docente para a personalização autônoma de critérios avaliativos em sistemas de IA.", style='List Bullet')

    add_heading("5 FUNDAMENTAÇÃO TEÓRICA")
    doc.add_paragraph(
        "A fundamentação teórica desta pesquisa converge três campos distintos. No eixo filosófico, recorre-se aos conceitos de Deleuze e Guattari "
        "sobre o pensamento geofilosófico, integrando as contribuições de Santos (2006) acerca do território usado. No eixo pedagógico, a base sustenta-se "
        "na perspectiva da Politecnia (SAVIANI, 2007) e no suporte da Neuropsicopedagogia através da Taxonomia de Bloom e do modelo SOLO. Por fim, "
        "no eixo tecnológico, adota-se o paradigma de Agentic Design Patterns (NG, 2024), que defende a superioridade de sistemas multiagentes especialistas."
    )

    add_heading("6 METODOLOGIA")
    doc.add_paragraph(
        "A pesquisa adota o procedimento da Design Science Research (DSR). O artefato consiste em um protótipo desenvolvido em Python com interface via Streamlit. "
        "A arquitetura prevê o uso de bancos de dados vetoriais (ChromaDB) para o suporte de um RAG temporal, bancos documentais (MongoDB) para a governança "
        "de prompts e unidades de processamento de baixa latência (LPUs). O fluxo baseia-se no ciclo agêntico de 'Caixa de Vidro Rastreável', onde o orquestrador "
        "coordena agentes especialistas focados nas dimensões cognitiva, práxica e territorial."
    )

    add_heading("7 CRONOGRAMA")
    doc.add_paragraph("[Inserir tabela de cronograma para 24 meses].")

    doc.add_page_break()

    # --- REFERÊNCIAS (NORMA ABNT COM NEGRITO APENAS NO TÍTULO) ---
    add_heading("8 REFERÊNCIAS")
    
    # Lista de tuplos (Parte em Negrito, Parte Normal)
    refs = [
        ("DELEUZE, Gilles; GUATTARI, Félix. ", "**O que é a filosofia?** Tradução de Bento Prado Júnior e Alberto Alonso Muñoz. 3. ed. Rio de Janeiro: Editora 34, 2010."),
        ("FERNANDES, Paulo Irineu Barreto. ", "**Introdução ao geofilosofar**: a hospitalidade como um novo agir. Uberlândia: [s. n.], 2023. (Tese de Doutorado)."),
        ("NG, Andrew. ", "**Agentic Design Patterns**: Part 1. [S. l.]: DeepLearning.AI, 20 mar. 2024. Disponível em: https://www.deeplearning.ai/the-batch/agentic-design-patterns-part-1/. Acesso em: 10 abr. 2026."),
        ("SANTOS, Milton. ", "**A natureza do espaço**: técnica e tempo, razão e emoção. 4. ed. 2. reimpr. São Paulo: EDUSP, 2006."),
        ("SAVIANI, Dermeval. ", "**A pedagogia histórico-crítica, as lutas de classe e a educação profissional**. Revista Trabalho, Educação e Saúde, Rio de Janeiro, v. 5, n. 2, p. 199-246, jul./out. 2007.")
    ]

    for autor, resto in sorted(refs):
        p = doc.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(12)
        
        p.add_run(autor)
        
        # Lógica para negritar apenas o título (até o primeiro ponto ou dois pontos)
        partes = resto.split("**")
        if len(partes) > 1:
            p.add_run(partes).bold = True
            p.add_run(partes)
        else:
            p.add_run(resto)

    doc.save("Pre_Projeto_Mestrado_IFTM_2026.docx")
    print("Ficheiro gerado com sucesso!")

if __name__ == "__main__":
    gerar_docx_mestrado()
