# /home/neirivon/SINAPSE2.0/sinapsebr_rubrica/scripts/Apresentacao.py
# --------------------------------------------------------------------------------------
# NOME DO SCRIPT: Apresentacao.py
# DESCRIÇÃO: Página inicial (Home) do protótipo SINAPSE-BR IA.
#            Apresenta o resumo estruturado do TCC, equipe e opções de exportação.
#
# FUNCIONALIDADES:
#   1. Resumo Expandido: Problema, Objetivos, Metodologia e Resultados.
#   2. Perfis: Cards visuais do Orientando e Orientadora.
#   3. Exportação: Geração dinâmica de DOCX e PDF com o conteúdo atualizado.
#   4. Navegação: Ponto de partida para as ferramentas (Rubrica, Timeline, Dados).
#
# AUTOR: Neirivon Elias Cardoso
# PROJETO: Rubrica SINAPSE-BR IA — Sistema Integrado Neuropsicopedagógico
# TCC: Pós-Graduação em Docência para a EPT (IFTM)
# DATA: 12/01/2026 (Atualizado conforme versão final do TCC)
# --------------------------------------------------------------------------------------

import os
from pathlib import Path
from io import BytesIO
import base64
from PIL import Image, ImageDraw
import streamlit as st
import io 

# --- IMPORTS PARA EXPORTAÇÃO ---
from docx import Document 
from fpdf import FPDF 
# -------------------------------------

# ---------------------------------
# Config da página
# ---------------------------------
st.set_page_config(
    page_title="SINAPSE-BR IA — Apresentação TCC",
    page_icon="🧠",
    layout="wide",
)

# ---------------------------------
# Localizador robusto de assets
# ---------------------------------
THIS = Path(__file__).resolve()

def find_project_root(start: Path, marker_folder: str = "assets") -> Path:
    """
    Sobe diretórios até encontrar uma pasta 'assets' (marcador do projeto).
    """
    p = start
    for _ in range(6): 
        if (p / marker_folder).exists():
            return p
        if p.parent == p:
            break
        p = p.parent
    return start

PROJECT_ROOT = find_project_root(THIS.parent)
ASSETS_DIR    = PROJECT_ROOT / "assets"
IMG_DIR       = ASSETS_DIR / "imagens"
LOGO_DIR      = ASSETS_DIR / "logos"

NEIRIVON_IMG      = IMG_DIR / "neirivon.png"
ORIENTADORA_IMG   = IMG_DIR / "Orientadora.png"
LOGO_IFTM         = LOGO_DIR / "IFTM_360.png"
LOGO_SINAPSE      = LOGO_DIR / "sinapse.png"

# ---------------------------------
# Utilitários de imagem / HTML
# ---------------------------------

def img_circular_b64(path: Path) -> str:
    """Converte imagem para Base64 e aplica máscara circular via PIL."""
    try:
        if not path.exists():
            return ""
        img = Image.open(path).convert("RGBA")
        mask = Image.new('L', img.size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, img.size[0], img.size[1]), fill=255)
        
        output_buffer = BytesIO()
        img.putalpha(mask)
        img.save(output_buffer, format="PNG")
        output_buffer.seek(0)
        
        base64_img = base64.b64encode(output_buffer.read()).decode()
        return f"data:image/png;base64,{base64_img}"
    except Exception as e:
        return ""

def tag_html_profile_content(base64_img: str, name: str, caption: str):
    """Gera o bloco HTML do perfil."""
    html_content = f"""
    <div class="profile-container">
        <img class="profile-image" src="{base64_img}" alt="{name}">
        <div class="profile-details">
            <h3>{name}</h3>
            <p>{caption}</p>
        </div>
    </div>
    """
    return html_content

def safe_image(path: Path, *, width: int | None = None, caption: str | None = None):
    try:
        if path.exists():
            st.image(str(path), width=width, caption=caption)
        else:
            st.warning(f"Imagem não encontrada: `{path.as_posix()}`") 
    except Exception as e:
        st.error(f"Erro ao carregar imagem: {e}")

# ---------------------------------
# FUNÇÕES DE GERAÇÃO DE ARQUIVOS (CONTEÚDO ATUALIZADO)
# ---------------------------------

def get_apresentacao_content():
    """
    Retorna o conteúdo COMPLETO da página principal atualizado com o TCC final.
    """
    content = []
    
    content.append("# 🧠 SINAPSE-BR IA")
    content.append("## Sistema Integrado Neuropsicopedagógico de Avaliação e Práticas da Sinergia Educacional Para a EPT")
    content.append("---")
    
    # 1. Autoria
    content.append("## 🧑‍🎓 Autoria e Orientação")
    content.append("**Orientando:** Neirivon Elias Cardoso")
    content.append("**Orientadora:** Profa. Dra. Thays Martins Vital da Silva")
    content.append("**Instituição:** IFTM - Campus Avançado Uberaba Parque Tecnológico")
    content.append("**Curso:** Pós-Graduação Lato Sensu em Docência para a EPT")
    content.append("---")

    # 2. Núcleo
    content.append("## 📚 Núcleo da Pesquisa (TEMA, PROBLEMA e DELIMITAÇÃO)")
    content.append("### TEMA")
    content.append("> Desenvolvimento de rubrica educacional ampliada para avaliação formativa na EPT, integrando Neuropsicopedagogia, DUA e Inteligência Artificial.")
    
    content.append("### PROBLEMA DE PESQUISA")
    content.append("`Como integrar princípios da Neuropsicopedagogia, das Taxonomias Cognitivas (Bloom/SOLO), do Desenho Universal para a Aprendizagem (DUA) e da equidade socioterritorial em uma rubrica formativa aplicável à Educação Profissional e Tecnológica?`")
    
    content.append("### DELIMITAÇÃO")
    # Texto atualizado para incluir a memória institucional nos documentos gerados
    content.append("Construção teórico-propositiva da **Rubrica SINAPSE-BR IA** para a Rede Federal, com recorte territorial e analítico no **Triângulo Mineiro e Alto Paranaíba (TMAP)**, utilizando dados do SISTEC, Censo Escolar e o **resgate da memória institucional**.")
    content.append("---")
    
    # 3. Introdução
    content.append("## 1. Introdução e Justificativa ✍️")
    content.append("A avaliação na EPT enfrenta o desafio de superar a simples verificação de tarefas técnicas. A **SINAPSE-BR IA** propõe um modelo que considera **como o cérebro aprende** (Neurociência) e **onde o aluno está** (Território).")
    
    content.append("### 🎯 Objetivo Geral")
    content.append("Analisar e propor uma rubrica educacional ampliada — SINAPSE-BR IA — fundamentada em referenciais de Neuropsicopedagogia, DUA e equidade territorial, visando aprimorar práticas avaliativas no contexto do TMAP.")
    
    content.append("### 🎯 Objetivos Específicos")
    content.append("""
* **1.** Analisar referenciais teóricos: Neuropsicopedagogia, Taxonomias e modelos de avaliação da EPT.
* **2.** Comparar estruturas de rubricas nacionais e internacionais para identificar lacunas.
* **3.** Propor a estrutura da Rubrica SINAPSE-BR IA e desenvolver o protótipo funcional.
    """)
    content.append("---")
    
    # 4. Fundamentação
    content.append("## 2. Fundamentação Teórica 🏗️")
    content.append("A pesquisa articula quatro eixos:")
    content.append("""
1. **EPT e Trabalho:** Saviani, Ramos, Frigotto e Ciavatta. O trabalho como princípio educativo e a formação integral.
2. **Neuropsicopedagogia:** Cosenza & Guerra, Piaget, Vygotsky. Foco nas funções executivas e neuroplasticidade.
3. **Inclusão e DUA:** Desenho Universal para a Aprendizagem (CAST) para eliminar barreiras avaliativas.
4. **Territorialização:** Uso de dados do SISTEC e Censo Escolar/INEP para contextualizar a oferta no TMAP.
    """)
    content.append("---")
    
    # 5. Metodologia
    content.append("## 3. Metodologia 🧪")
    content.append("**Natureza:** Pesquisa Teórico-Propositiva, Qualitativa e Descritiva.")
    content.append("### Procedimentos:")
    content.append("""
* **Levantamento Bibliográfico:** Revisão narrativa em bases de dados e repositórios.
* **Análise Documental:** DCNs, BNCC, Matrizes do SAEB e Relatórios de Gestão.
* **Construção do Artefato:** Desenvolvimento do software SINAPSE em Python/Streamlit.
    """)
    content.append("---")

    # 6. Produto
    content.append("## 4. Produto Educacional: SINAPSE-BR IA 🖥️")
    content.append("O software entrega:")
    content.append("""
- **Rubrica Interpretativa:** Avaliação baseada em níveis de processamento cognitivo.
- **Painel Territorial:** Mapa da oferta EPT no TMAP (2017-2024).
- **Relatórios Automatizados:** Apoio à gestão e à docência.
    """)
    content.append("---")

    content.append("## 5. Considerações Finais ✅")
    content.append("A SINAPSE-BR IA não é apenas uma ferramenta de notas, mas um sistema de apoio à decisão que integra **dados** e **pedagogia**, promovendo uma avaliação mais justa e humana na Rede Federal.")
    
    return "\n\n".join(content)

def generate_docx(content_markdown: str) -> BytesIO:
    document = Document()
    document.add_heading("SINAPSE-BR IA - Apresentação TCC", 0)
    
    for line in content_markdown.split('\n\n'):
        line = line.strip()
        if not line: continue
        
        if line.startswith('## '):
            document.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            document.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('>'):
            p = document.add_paragraph(line.replace('>', '').strip())
            p.style = 'Intense Quote'
        elif line.startswith('`'):
            document.add_paragraph(line.replace('`', '').strip()) 
        elif '*' in line and line.strip().startswith('*'):
            for item in line.split('\n'):
                if item.strip():
                    document.add_paragraph(item.replace('*', '').strip(), style='List Bullet')
        else:
            document.add_paragraph(line)

    doc_buffer = BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def generate_pdf(content_markdown: str) -> BytesIO:
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 12)
            self.cell(0, 10, 'SINAPSE-BR IA - Apresentação TCC', 0, 1, 'C')
            self.ln(5)
        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Página {self.page_no()}/{{nb}}', 0, 0, 'C')

    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    for line in content_markdown.split('\n\n'):
        line = line.strip()
        if not line:
            pdf.ln(2)
            continue
            
        if line.startswith('##'):
            level = line.count('#')
            text = line.replace('#', '').strip()
            font_size = 14 if level == 1 else 12
            pdf.set_font('Arial', 'B', font_size)
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 8, text.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(2)
        elif line.startswith('>'):
            pdf.set_font('Arial', 'I', 10)
            pdf.set_text_color(80, 80, 80)
            text = line.replace('>', '').strip()
            pdf.multi_cell(0, 6, text.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(2)
        elif line.startswith('`'):
            pdf.set_font('Courier', '', 10)
            pdf.set_text_color(0, 0, 128)
            text = line.replace('`', '').strip()
            pdf.multi_cell(0, 6, text.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(2)
        else:
            pdf.set_font('Arial', '', 10)
            pdf.set_text_color(0, 0, 0)
            text = line.replace('*', '').strip()
            pdf.multi_cell(0, 6, text.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(1)

    pdf_output = pdf.output(dest='S')
    return BytesIO(pdf_output)

# ---------------------------------
# SIDEBAR
# ---------------------------------
with st.sidebar:
    safe_image(LOGO_IFTM, width=130) 
    st.markdown("---")
    st.markdown("### 📑 Navegação TCC")
    st.caption("Apresentação do Trabalho de Conclusão de Curso.")
    st.markdown("---")
    safe_image(LOGO_SINAPSE, width=200)
    st.caption("SINAPSE-BR • Sistema Integrado Neuropsicopedagógico")
    st.markdown("---")
    
# ---------------------------------
# PÁGINA PRINCIPAL
# ---------------------------------

# Cabeçalho na Tela
st.title("🧠 SINAPSE-BR IA")
st.subheader("Sistema Integrado Neuropsicopedagógico de Avaliação e Práticas da Sinergia Educacional Para a EPT")
st.markdown("**Curso:** Pós-Graduação em Docência para a EPT — IFTM")

col_side, col_main = st.columns([0.25, 0.75])

with col_main:
    
    # CSS Injetado
    st.markdown("""
    <style>
        .profile-container {
            display: flex; align-items: center; padding: 15px;
            background-color: #f8f9fa; border-radius: 10px;
            margin-bottom: 15px; border: 1px solid #e0e0e0;
        }
        .profile-image {
            width: 120px; height: 120px; border-radius: 50%;
            object-fit: cover; margin-right: 20px;
            border: 3px solid #4CAF50;
        }
        .profile-details h3 { margin: 0 0 5px 0; color: #2c3e50; font-size: 1.3em;}
        .profile-details p { margin: 0; color: #7f8c8d; font-size: 0.95em;}
    </style>
    """, unsafe_allow_html=True)

    # Perfis
    st.markdown("### 🧑‍🎓 Autoria e Orientação")
    if NEIRIVON_IMG.exists():
        st.markdown(tag_html_profile_content(
            img_circular_b64(NEIRIVON_IMG), 
            "Neirivon Elias Cardoso", 
            "Especialista em Docência EPT"
        ), unsafe_allow_html=True)
        
    if ORIENTADORA_IMG.exists():
        st.markdown(tag_html_profile_content(
            img_circular_b64(ORIENTADORA_IMG), 
            "Dra. Thays Martins Vital da Silva", 
            "Orientadora do TCC"
        ), unsafe_allow_html=True)

    st.divider()

    # Expander: O Coração da Pesquisa
    with st.expander("📚 Núcleo da Pesquisa (Problema e Delimitação)", expanded=True):
        st.info("**TEMA:** Desenvolvimento de rubrica educacional ampliada para EPT com base em Neuropsicopedagogia e IA.")
        st.markdown("**PROBLEMA:**")
        st.code("Como integrar princípios da Neuropsicopedagogia, das Taxonomias Cognitivas (Bloom/SOLO), do DUA e da equidade socioterritorial em uma rubrica formativa aplicável à EPT?", language="text")
        
        # --- ATUALIZAÇÃO DA DELIMITAÇÃO PARA A BANCA ---
        st.markdown("**DELIMITAÇÃO:**")
        st.markdown("""
        Foco na Rede Federal (IFTM), recorte territorial TMAP, utilizando:
        * Dados quantitativos do SISTEC e Censo Escolar;
        * **Resgate da memória institucional** e dados qualitativos do território.
        """)
        # -----------------------------------------------

    # Abas de Conteúdo
    tab1, tab2, tab3, tab4 = st.tabs(["1. Justificativa", "2. Teoria (EPT/Neuro)", "3. Metodologia", "4. Produto SINAPSE"])

    with tab1:
        st.markdown("### 🚀 Por que SINAPSE-BR?")
        st.write("A avaliação na EPT muitas vezes reproduz a dualidade estrutural. A SINAPSE-BR busca superar isso através da **Neuropsicopedagogia**, oferecendo uma rubrica que considera *como* o aluno aprende.")
        st.markdown("#### Objetivos")
        st.success("Geral: Propor a Rubrica SINAPSE-BR IA como instrumento de equidade e avaliação formativa.")

    with tab2:
        st.markdown("### 🏗️ Fundamentação Teórica")
        c1, c2 = st.columns(2)
        with c1:
            st.markdown("#### EPT & Sociedade")
            st.caption("Saviani, Frigotto, Ramos")
            st.write("Trabalho como princípio educativo e formação omnilateral.")
        with c2:
            st.markdown("#### Neurociência & Cognição")
            st.caption("Cosenza, Guerra, Piaget")
            st.write("Funções executivas e níveis taxonômicos (Bloom/SOLO).")

    with tab3:
        st.markdown("### 🧪 Percurso Metodológico")
        st.write("Pesquisa **Qualitativa, Teórico-Propositiva**.")
        st.markdown("""
        1. **Levantamento Bibliográfico:** Bases SciELO, Repositórios.
        2. **Análise Documental:** DCNs, SAEB, Relatórios SISTEC.
        3. **Engenharia Didática:** Construção do artefato em Python (Streamlit).
        """)

    with tab4:
        st.markdown("### 🖥️ O Produto")
        st.write("Aplicação Web que une avaliação e território.")
        st.json({
            "Componente 1": "Rubrica Interpretativa",
            "Componente 2": "Painel de Dados Territoriais",
            "Fim": "Apoio à decisão docente"
        })

    st.divider()

    # Área de Download
    st.subheader("⬇️ Baixar Apresentação")
    
    # Gerar conteúdo atualizado
    md_content = get_apresentacao_content()
    
    col_d1, col_d2, col_d3 = st.columns([1, 1, 2])
    with col_d1:
        st.download_button(
            "📄 Baixar DOCX", 
            data=generate_docx(md_content), 
            file_name="SINAPSE_Apresentacao_TCC.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col_d2:
        try:
            st.download_button(
                "📑 Baixar PDF", 
                data=generate_pdf(md_content), 
                file_name="SINAPSE_Apresentacao_TCC.pdf",
                mime="application/pdf"
            )
        except Exception as e:
            st.error("Erro PDF (instale fpdf2)")

    st.caption(f"Sistema rodando a partir de: `{PROJECT_ROOT}`")
