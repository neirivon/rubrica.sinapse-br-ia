# /home/neirivon/SINAPSE2.0/sinapsebr_rubrica/scripts/Apresentacao.py
import os
from pathlib import Path
from io import BytesIO
import base64
from PIL import Image, ImageDraw
import streamlit as st
import io 

# --- NOVOS IMPORTS PARA EXPORTAÇÃO ---
from docx import Document 
from fpdf import FPDF 
# -------------------------------------

# ---------------------------------
# Config da página
# ---------------------------------
st.set_page_config(
    page_title="SINAPSE-BR IA — Apresentação",
    page_icon="🧠",
    layout="wide",
)

# ---------------------------------
# Localizador robusto de assets
# ---------------------------------
THIS = Path(__file__).resolve()

def find_project_root(start: Path, marker_folder: str = "assets") -> Path:
    """
    Sobe diretórios até encontrar uma pasta 'assets' (marcador do projeto).
    Retorna o diretório onde 'assets' foi encontrado; se não encontrar,
    retorna o diretório do arquivo atual.
    """
    p = start
    for _ in range(6): 
        if (p / marker_folder).exists():
            return p
        if p.parent == p:
            break
        p = p.parent
    return start

# Se o script foi movido para a raiz, esta linha pode precisar ser ajustada
# para: PROJECT_ROOT = find_project_root(THIS)
PROJECT_ROOT = find_project_root(THIS.parent)
ASSETS_DIR    = PROJECT_ROOT / "assets"
IMG_DIR       = ASSETS_DIR / "imagens"
LOGO_DIR      = ASSETS_DIR / "logos"

NEIRIVON_IMG      = IMG_DIR / "neirivon.png"
ORIENTADORA_IMG   = IMG_DIR / "Orientadora.png"
LOGO_IFTM         = LOGO_DIR / "IFTM_360.png"
LOGO_SINAPSE      = LOGO_DIR / "sinapse.png"

# ---------------------------------
# Utilitários de imagem / HTML
# ---------------------------------

def img_circular_b64(path: Path) -> str:
    """Converte imagem para Base64 e a transforma em círculo usando PIL."""
    try:
        if not path.exists():
            return ""

        img = Image.open(path).convert("RGBA")
        mask = Image.new('L', img.size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, img.size[0], img.size[1]), fill=255)
        
        output_buffer = BytesIO()
        img.putalpha(mask)
        img.save(output_buffer, format="PNG")
        output_buffer.seek(0)
        
        base64_img = base64.b64encode(output_buffer.read()).decode()
        return f"data:image/png;base64,{base64_img}"
    except Exception as e:
        return ""

def tag_html_profile_content(base64_img: str, name: str, caption: str):
    """Gera APENAS o bloco HTML do perfil (DIV), confiando no CSS injetado no topo."""
    html_content = f"""
    <div class="profile-container">
        <img class="profile-image" src="{base64_img}" alt="{name}">
        <div class="profile-details">
            <h3>{name}</h3>
            <p>{caption}</p>
        </div>
    </div>
    """
    return html_content

def safe_image(path: Path, *, width: int | None = None, caption: str | None = None):
    """Exibe imagem se existir; caso contrário, mostra um aviso discreto."""
    try:
        if path.exists():
            st.image(str(path), width=width, caption=caption)
        else:
            st.warning(f"Imagem não encontrada: `{path.as_posix()}`") 
    except Exception as e:
        st.error(f"Não foi possível carregar a imagem `{path.name}`. Detalhe: {e}")

# ---------------------------------
# FUNÇÕES DE GERAÇÃO DE ARQUIVOS
# ---------------------------------

def get_apresentacao_content():
    """
    Retorna o conteúdo COMPLETO da página principal como uma string Markdown.
    ATUALIZADO: Inclui SAEB na Seção 4 e restaura Seções 5 e 6.
    """
    content = []
    
    content.append("# 🧠 SINAPSE-BR IA — Rubrica Avaliativa Interpretativa para a EPT")
    content.append("---")
    
    # 1. Orientando
    content.append("## 🧑‍🎓 Orientando")
    content.append("### Neirivon Elias Cardoso")
    content.append("_Orientando do TCC_")
    content.append("---")

    # 2. Orientadora
    content.append("## 👩‍🏫 Orientadora")
    content.append("### Dra. Professora Thays Martins Vital da Silva")
    content.append("_Orientadora do TCC_")
    content.append("---")

    # 3. Núcleo da Proposta
    content.append("## 📚 Núcleo da Proposta (TEMA, PROBLEMA e DELIMITAÇÃO)")
    content.append("### TEMA")
    content.append(
        "> Desenvolvimento de uma rubrica educacional ampliada para avaliação formativa na Educação Profissional e Tecnológica (EPT), "
        "integrando referenciais da Neuropsicopedagogia, Taxonomias Cognitivas e Inteligência Artificial."
    )
    content.append("### PROBLEMA DE PESQUISA")
    content.append(
        "`Como integrar princípios da Neuropsicopedagogia, das Taxonomias Cognitivas e da equidade socio-territorial em uma rubrica formativa aplicável à Educação Profissional e Tecnológica?`"
    )
    content.append("### DELIMITAÇÃO DO TEMA")
    content.append(
        "O estudo concentra-se na construção teórico-propositiva da **Rubrica SINAPSE-BR IA**, concebida para qualificar práticas avaliativas na Rede Federal de EPT, com ênfase no recorte territorial do **Triângulo Mineiro e Alto Paranaíba (TMAP)**."
    )
    content.append("---")
    
    # SEÇÃO 1: INTRODUÇÃO
    content.append("## 1. Introdução & Estratégia da Pesquisa ✍️")
    content.append("### ✅ Justificativa")
    
    content.append(
        "A avaliação na Educação Profissional e Tecnológica apresenta desafios relacionados à clareza dos critérios e à equidade territorial. As rubricas tradicionais muitas vezes não contemplam as especificidades cognitivas dos estudantes da EPT.\n\n"
        "A criação da Rubrica SINAPSE-BR IA busca preencher essa lacuna, fundamentando-se na convergência entre a **Neurociência da Aprendizagem** (Cosenza & Guerra) e a **Neuropsicopedagogia** (Chupil; Lopes), criando um instrumento coerente, formativo e sensível às realidades do TMAP."
    )

    content.append("### 🎯 Objetivo Geral")
    content.append(
        "Desenvolver uma rubrica educacional ampliada — denominada **SINAPSE-BR IA** — com vistas a aprimorar as práticas avaliativas na Educação Profissional e Tecnológica e favorecer trajetórias formativas mais justas no contexto do **Triângulo Mineiro e Alto Paranaíba (TMAP)**."
    )
    
    content.append("### 🎯 Objetivos Específicos")
    content.append("""
* **1.** Analisar os referenciais teóricos da Neuropsicopedagogia, Taxonomias Cognitivas e modelos de avaliação da EPT.
* **2.** Comparar estruturas de rubricas nacionais e internacionais para identificar lacunas avaliativas.
* **3.** Propor a estrutura final da Rubrica SINAPSE-BR IA, articulando fundamentos pedagógicos e neurocientíficos.
    """)
    content.append("---")
    
    # SEÇÃO 2: FUNDAMENTAÇÃO TEÓRICA
    content.append("## 2. Fundamentação Teórica 📚")
    content.append("A proposta fundamenta-se na articulação de quatro eixos principais:")
    content.append(
        "**1. Neuropsicopedagogia e Neurociência:** Baseada em Cosenza & Guerra (2011) e Dehaene, abordando como o cérebro aprende (neuroplasticidade, atenção) e o campo de atuação profissional (Chupil et al., 2018).\n\n"
        "**2. Taxonomias e Cognição:** Utiliza Piaget e Vygotsky para o desenvolvimento, e a Taxonomia de Bloom revisada (Anderson & Krathwohl) para estruturar os níveis de complexidade.\n\n"
        "**3. Princípios Transversais (Inclusão):** Adota o Desenho Universal para a Aprendizagem (DUA/CAST) e as Metodologias Ativas como contexto para garantir a acessibilidade avaliativa.\n\n"
        "**4. Territorialização:** Compreensão da oferta educacional baseada nas definições oficiais do IBGE e dados do Censo Escolar."
    )
    content.append("---")
    
    # SEÇÃO 3: METODOLOGIA
    content.append("## 3. Metodologia 🧪")
    content.append("**Tipo de pesquisa:** teórico-propositiva, qualitativa e descritiva, com desenvolvimento de protótipo digital.")
    
    content.append("### 1. Revisão de Literatura")
    content.append("O embasamento teórico estrutura-se nos seguintes núcleos:")
    content.append("""
* **Núcleo Neurocientífico:** Cosenza & Guerra (2011), Dehaene (Science of Education).
* **Núcleo Neuropsicopedagógico:** Obras brasileiras de referência como Chupil, Souza & Schneider (2018) e Lopes (2020).
* **Núcleo Tecnológico (IA):** Russell & Norvig (2022) para fundamentação de agentes inteligentes e Nicolelis/Seung para a metáfora de redes.
* **Núcleo Pedagógico:** Piaget, Vygotsky, Flavell (Metacognição) e Hoffmann (Avaliação Mediadora).
""")
    
    content.append("### 2. Análise Documental")
    content.append("""
* Diretrizes Curriculares Nacionais da EPT.
* Matrizes do SAEB e PISA/OCDE.
* Microdados do Censo Escolar e SISTEC.
""")
    
    content.append("### 3. Construção do Artefato (SINAPSE-BR IA)")
    content.append(
        "Desenvolvimento do protótipo em Python/Streamlit, integrando os dados territoriais com a lógica da rubrica fundamentada."
    )
    content.append("---")

    # SEÇÃO 4: PRODUTO EDUCACIONAL (ATUALIZADA COM SAEB)
    content.append("## 4. Produto Educacional 🖥️")
    content.append(
        "Aplicativo **SINAPSE-BR IA** em **Streamlit** com:\n"
        "- Menu lateral (logos IFTM e SINAPSE-BR).\n"
        "- Página de **Apresentação** (orientando + orientadora).\n"
        "- Páginas territoriais (2017-2024) com dados reais.\n"
        "- Execução local e nuvem.\n"
        "- **Fontes de Dados:** SISTEC, INEP (Censo Escolar) e **SAEB**."
    )
    content.append("---")
    
    # SEÇÃO 5: RESULTADOS ESPERADOS (RESTAURADA)
    content.append("## 5. Resultados Esperados 🎯")
    content.append(
        "- Visualizações confiáveis da **rede EPT** no **TMAP**.\n"
        "- Identificação de **lacunas regionais** de oferta e infraestrutura.\n"
        "- Apoio ao docente na **avaliação formativa** baseada em evidências.\n"
        "- Ferramenta para gestores analisarem **equidade territorial** (INSE x SAEB).\n"
        "- Base para **instrumentos avaliativos personalizados**."
    )
    content.append("---")

    # SEÇÃO 6: DISCUSSÃO (RESTAURADA)
    content.append("## 6. Discussão 💬")
    content.append(
        "A integração entre **dados abertos** (SAEB/Censo), **pedagogia** e **territorialização** fortalece políticas públicas educacionais. "
        "A proposta demonstra que a Neuropsicopedagogia pode qualificar a leitura de dados educacionais, transformando estatísticas em estratégias pedagógicas. "
        "A análise revela que a 'invisibilidade' de dados em áreas rurais (N/A) exige rubricas locais sensíveis ao contexto."
    )
    content.append("---")
    
    # SEÇÃO 7: CONSIDERAÇÕES FINAIS
    content.append("## 7. Considerações Finais ✅")
    content.append(
        "A **Rubrica SINAPSE-BR IA** avança na integração entre **dados educacionais** e **fundamentos neuropsicopedagógicos**. "
        "O trabalho futuro prevê a validação empírica da rubrica com docentes da Rede Federal, consolidando o SINAPSE como um Recurso Educacional Digital (RED) de apoio à decisão docente."
    )
    
    return "\n\n".join(content)


def generate_docx(content_markdown: str) -> BytesIO:
    """
    Gera o arquivo DOCX em memória a partir do conteúdo Markdown.
    """
    document = Document()
    document.add_heading("SINAPSE-BR IA - Apresentação", 0)
    
    for line in content_markdown.split('\n\n'):
        if line.startswith('#'):
            level = line.count('#')
            text = line.lstrip('#').strip()
            if text:
                document.add_heading(text, level=min(level, 4))
        elif line.startswith('>'):
            document.add_paragraph(line.lstrip('>').strip(), style='Intense Quote')
        elif line.startswith('`'):
            document.add_paragraph(line.strip('`').strip()) 
        elif line.startswith('*'):
            document.add_paragraph(line.strip('*').strip(), style='List Bullet')
        elif line.strip():
            document.add_paragraph(line.strip())

    doc_buffer = BytesIO()
    document.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer

def generate_pdf(content_markdown: str) -> BytesIO:
    """
    Gera o arquivo PDF em memória usando fpdf2 a partir do conteúdo Markdown.
    """
    class PDF(FPDF):
        def header(self):
            self.set_font('Arial', 'B', 15)
            self.cell(0, 10, 'SINAPSE-BR IA - Apresentação'.encode('latin-1', 'replace').decode('latin-1'), 0, 1, 'C')
            self.ln(5)

        def footer(self):
            self.set_y(-15)
            self.set_font('Arial', 'I', 8)
            self.cell(0, 10, f'Página {self.page_no()}/{{nb}}'.encode('latin-1', 'replace').decode('latin-1'), 0, 0, 'C')

    pdf = PDF()
    pdf.alias_nb_pages()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15) 
    
    for line in content_markdown.split('\n\n'):
        line = line.strip()
        if not line:
            pdf.ln(3) 
            continue
            
        if line.startswith('#'):
            level = line.count('#')
            text = line.lstrip('#').strip()
            if text:
                font_size = max(10, 16 - level * 2) 
                pdf.set_font('Arial', 'B', font_size)
                pdf.multi_cell(0, 8, text.encode('latin-1', 'replace').decode('latin-1'))
                pdf.ln(2)
        
        else:
            pdf.set_font('Arial', '', 10)
            pdf.set_text_color(0, 0, 0) 
            text_to_render = line.replace('\n', ' ')
            line_height = 5
            
            if text_to_render.startswith('>'):
                 pdf.set_text_color(100, 100, 100)
                 pdf.set_font('Arial', 'I', 10)
                 text_to_render = text_to_render.lstrip('>').strip()
                 line_height = 6
            elif text_to_render.startswith('`'):
                 pdf.set_text_color(0, 0, 128)
                 text_to_render = text_to_render.strip('`').strip()
                 line_height = 6

            pdf.write(line_height, text_to_render.encode('latin-1', 'replace').decode('latin-1'))
            pdf.ln(line_height) 

    pdf_output = pdf.output(dest='S')
    return BytesIO(pdf_output)


# ---------------------------------
# SIDEBAR — IFTM (topo), marcador da página, SINAPSE (rodapé)
# ---------------------------------
with st.sidebar:
    safe_image(LOGO_IFTM, width=130) 
    st.markdown("---")
    st.markdown("### 📑 Apresentação")
    st.caption("Visão geral do TCC, equipe e referências iniciais.")
    
    st.markdown("---")
    safe_image(LOGO_SINAPSE, width=200)
    st.caption("SINAPSE-BR • Sistema Integrado Neuropsicopedagógico")
    st.markdown("---")
    
# ---------------------------------
# Cabeçalho
# ---------------------------------
st.header("🧠 SINAPSE-BR IA — Rubrica Avaliativa Interpretativa para a EPT")

# ---------------------------------
# Layout (conteúdo)
# ---------------------------------
col_side, col_main = st.columns([0.18, 0.82])

with col_main:
    
    # ---------------------------------
    # INJEÇÃO DO CSS GLOBALMENTE NO TOPO DO col_main
    # ---------------------------------
    global_css_block = """
    <style>
        .profile-container {
            display: flex;
            align-items: center;
            padding: 10px 10px;
            margin-bottom: 10px;
            border-radius: 10px;
            transition: background-color 0.3s ease;
        }
        .profile-container:hover {
            background-color: #e0e0e0; /* Efeito mouseover */
        }
        .profile-image {
            width: 150px;
            height: 150px;
            border-radius: 50%;
            object-fit: cover;
            margin-right: 20px;
            box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        }
        .profile-details h3 {
            margin-top: 0;
            margin-bottom: 5px;
            font-size: 1.5em;
            color: #333;
        }
        .profile-details p {
            margin: 0;
            font-size: 0.9em;
            color: #666;
        }
    </style>
    """
    st.markdown(global_css_block, unsafe_allow_html=True)

    # ---------------------------------
    # -------- Orientando e Orientadora --------
    # ---------------------------------
    st.subheader("🧑‍🎓 Orientando")
    if NEIRIVON_IMG.exists():
        neirivon_b64 = img_circular_b64(NEIRIVON_IMG)
        profile_html = tag_html_profile_content(neirivon_b64, "Neirivon Elias Cardoso", "Orientando do TCC")
        st.markdown(profile_html, unsafe_allow_html=True)
    else:
        st.error(f"Imagem não encontrada: {NEIRIVON_IMG}")

    st.divider()

    st.subheader("👩‍🏫 Orientadora")
    if ORIENTADORA_IMG.exists():
        orientadora_b64 = img_circular_b64(ORIENTADORA_IMG)
        profile_html = tag_html_profile_content(orientadora_b64, "Dra. Professora Thays Martins Vital da Silva", "Orientadora do TCC")
        st.markdown(profile_html, unsafe_allow_html=True)
    else:
        st.error(f"Imagem não encontrada: {ORIENTADORA_IMG}")

    st.divider()
    
    # ---------------------------------
    # SEÇÃO 0: TEMA (Atualizado)
    # ---------------------------------
    with st.expander("📚 Núcleo da Proposta (TEMA, PROBLEMA e DELIMITAÇÃO)", expanded=False):
        st.markdown("### TEMA")
        st.info(
            "Desenvolvimento de uma rubrica educacional ampliada para avaliação formativa na Educação Profissional e Tecnológica (EPT), "
            "integrando referenciais da Neuropsicopedagogia, Taxonomias Cognitivas e Inteligência Artificial."
        )
        
        st.markdown("### PROBLEMA DE PESQUISA")
        st.code(
            "Como integrar princípios da Neuropsicopedagogia, das Taxonomias Cognitivas e da equidade socio-territorial em uma rubrica formativa aplicável à Educação Profissional e Tecnológica?",
            language="markdown"
        )

        st.markdown("### DELIMITAÇÃO DO TEMA")
        st.write(
            "O estudo concentra-se na construção teórico-propositiva da **Rubrica SINAPSE-BR IA**, concebida para qualificar práticas avaliativas na Rede Federal de EPT, com ênfase no recorte territorial do **Triângulo Mineiro e Alto Paranaíba (TMAP)**."
        )

    st.markdown("---")

    # ---------------------------------
    # SEÇÃO 1: INTRODUÇÃO
    # ---------------------------------
    st.markdown("## 1. Introdução & Estratégia da Pesquisa ✍️")
    
    tab_justificativa, tab_objetivos, tab_sinapse_ia = st.tabs([
        "✅ Justificativa", 
        "🎯 Objetivos", 
        "🧠 Visão Geral do SINAPSE-BR IA"
    ])

    with tab_justificativa:
        st.markdown("### Por Que SINAPSE-BR IA?")
        st.write(
            "A criação da Rubrica SINAPSE-BR IA busca preencher as lacunas avaliativas da EPT, fundamentando-se na **Neurociência da Aprendizagem** (Cosenza & Guerra; Dehaene) e na **Neuropsicopedagogia** (Chupil; Lopes), criando um instrumento coerente e formativo."
        )

    with tab_objetivos:
        st.markdown("### Objetivo Geral")
        st.info(
            "Desenvolver uma rubrica educacional ampliada — denominada **SINAPSE-BR IA** — com vistas a aprimorar as práticas avaliativas na Educação Profissional e Tecnológica e favorecer trajetórias formativas mais justas no contexto do **Triângulo Mineiro e Alto Paranaíba (TMAP)**."
        )
        st.markdown("### Objetivos Específicos")
        st.markdown("""
            * **1.** Analisar os referenciais teóricos da Neuropsicopedagogia, Taxonomias Cognitivas e modelos de avaliação da EPT.
            * **2.** Comparar estruturas de rubricas nacionais e internacionais para identificar lacunas avaliativas.
            * **3.** Propor a estrutura final da Rubrica SINAPSE-BR IA, articulando fundamentos pedagógicos e neurocientíficos.
        """)
        
    with tab_sinapse_ia:
        st.markdown("### Sobre o Protótipo SINAPSE-BR IA")
        st.write(
            "O presente TCC propõe a **Rubrica Educacional SINAPSE-BR IA**. O protótipo computacional permite visualizar a **oferta real da EPT** nos municípios do TMAP, integrando dados e pedagogia."
        )

    st.markdown("---")

    # ---------------------------------
    # SEÇÃO 2: FUNDAMENTAÇÃO TEÓRICA
    # ---------------------------------
    st.markdown("## 2. Fundamentação Teórica 📚")
    
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**1. Neuropsicopedagogia e Neurociência**")
        st.caption("Cosenza & Guerra (2011), Dehaene e Chupil et al. (2018). Foco na aprendizagem biológica e intervenção.")
        st.markdown("**2. Taxonomias e Cognição**")
        st.caption("Piaget, Vygotsky e Bloom (Anderson & Krathwohl). Foco no desenvolvimento e níveis de complexidade.")
    with col2:
        st.markdown("**3. Princípios Transversais (Inclusão)**")
        st.caption("Desenho Universal para a Aprendizagem (DUA) e Metodologias Ativas. Acessibilidade como contexto.")
        st.markdown("**4. Territorialização**")
        st.caption("Dados oficiais do IBGE e Censo Escolar para análise contextual.")

    st.markdown("---")

    # ---------------------------------
    # SEÇÃO 3: METODOLOGIA
    # ---------------------------------
    st.markdown("## 3. Metodologia 🧪")
    st.write("**Tipo de pesquisa:** teórico-propositiva, qualitativa e descritiva, com desenvolvimento de protótipo digital.")
    
    tab_revisao, tab_analise, tab_construcao = st.tabs([
        "1. Revisão de Literatura", 
        "2. Análise Documental", 
        "3. Construção Propositiva"
    ])
    
    with tab_revisao:
        st.markdown("### 1. Revisão de Literatura")
        st.write("A fundamentação teórica será estruturada em quatro eixos de convergência:")
        st.markdown("""
        * **Núcleo Neurocientífico:** Cosenza & Guerra (2011), Dehaene.
        * **Núcleo Neuropsicopedagógico:** Chupil, Souza & Schneider (2018), Lopes (2020).
        * **Núcleo Tecnológico (IA):** Russell & Norvig (2022) e Nicolelis/Seung.
        * **Núcleo Pedagógico:** Piaget, Vygotsky, Flavell e Hoffmann.
        """)

    with tab_analise:
        st.markdown("### 2. Análise Documental Comparativa")
        st.markdown("""
        Serão analisados documentos oficiais e modelos avaliativos, incluindo:
        * BNCC, SAEB, PISA/OCDE, Diretrizes da EPT.
        * DUA, rubricas nacionais e internacionais (Andrade; Brookhart; Mullinix; Moskal).
        * Materiais normativos da Rede Federal.
        """)
        st.caption("A análise busca identificar convergências, divergências e lacunas que justifiquem a necessidade de uma rubrica integradora adequada ao contexto da Educação Profissional e Tecnológica, especialmente no TMAP.")

    with tab_construcao:
        st.markdown("### 3. Construção Propositiva da Rubrica SINAPSE-BR IA")
        st.write(
            "Será elaborada a versão final da rubrica (dimensões, níveis e descritores), integrando fundamentos neurocientíficos, pedagógicos e socio-territoriais. A rubrica será organizada para favorecer práticas avaliativas formativas, inclusivas e alinhadas à realidade da EPT. Serão indicadas, ainda, possibilidades de aplicação prática futura no contexto educacional da região TMAP."
        )
        st.markdown("### 3.4 Fontes de Dados do Protótipo")
        st.markdown("""
        * **Relatório IPES Escolas (2020–2023)** — SISTEC
        * **Sistec Cursos Técnicos Ativos (12/09/2022)**
        * **Suplemento Cursos Técnicos 2024** — Censo Escolar/INEP
        """)
        st.write("Mapeamento IBGE, normalização de cabeçalhos e valores; **nenhum dado inventado/estimado**.")

    st.markdown("---")

    # ---------------------------------
    # SEÇÃO 4: PRODUTO EDUCACIONAL (ATUALIZADA)
    # ---------------------------------
    st.markdown("## 4. Produto Educacional 🖥️")
    st.write(
        "Aplicativo **SINAPSE-BR IA** em **Streamlit** com:\n"
        "- Menu lateral (logos IFTM e SINAPSE-BR).\n"
        "- Página de **Apresentação** (orientando + orientadora).\n"
        "- Páginas territoriais (2017-2024) com dados reais.\n"
        "- Execução local e nuvem.\n"
        "- **Fontes de Dados:** SISTEC, INEP (Censo Escolar) e **SAEB**."
    )

    st.markdown("---")

    # ---------------------------------
    # SEÇÃO 5: RESULTADOS ESPERADOS (RESTAURADA)
    # ---------------------------------
    st.markdown("## 5. Resultados Esperados 🎯")
    st.write(
        "- Visualizações confiáveis da **rede EPT** no **TMAP**.\n"
        "- Identificação de **lacunas regionais** de oferta e infraestrutura.\n"
        "- Apoio ao docente na **avaliação formativa** baseada em evidências.\n"
        "- Ferramenta para gestores analisarem **equidade territorial** (INSE x SAEB).\n"
        "- Base para **instrumentos avaliativos personalizados**."
    )

    st.markdown("---")

    # ---------------------------------
    # SEÇÃO 6: DISCUSSÃO (RESTAURADA)
    # ---------------------------------
    st.markdown("## 6. Discussão 💬")
    st.write(
        "A integração entre **dados abertos** (SAEB/Censo), **pedagogia** e **territorialização** fortalece políticas públicas educacionais. "
        "A proposta demonstra que a Neuropsicopedagogia pode qualificar a leitura de dados educacionais, transformando estatísticas em estratégias pedagógicas. "
        "A análise revela que a 'invisibilidade' de dados em áreas rurais (N/A) exige rubricas locais sensíveis ao contexto."
    )

    st.markdown("---")

    # ---------------------------------
    # SEÇÃO 7: CONSIDERAÇÕES FINAIS
    # ---------------------------------
    st.markdown("## 7. Considerações Finais ✅")
    st.write(
        "A **Rubrica SINAPSE-BR IA** avança na integração entre **dados educacionais** e **fundamentos neuropsicopedagógicos**. "
        "O trabalho futuro prevê a validação empírica da rubrica com docentes da Rede Federal, consolidando o SINAPSE como um Recurso Educacional Digital (RED) de apoio à decisão docente."
    )

    st.markdown("---")

    # --- BOTÕES DE DOWNLOAD MOVIDOS PARA O FINAL DA PÁGINA PRINCIPAL (col_main) ---
    st.subheader("⬇️ Baixar Apresentação")
    
    # Geração do conteúdo Markdown (já feita antes)
    apresentacao_markdown = get_apresentacao_content()
    
    col_docx, col_pdf, col_spacer = st.columns([0.2, 0.2, 0.6])

    with col_docx:
        # Geração DOCX
        docx_buffer = generate_docx(apresentacao_markdown)
        st.download_button(
            label="Baixar DOCX",
            data=docx_buffer,
            file_name="SINAPSE_BR_IA_Apresentacao.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Baixa a página principal no formato Microsoft Word (.docx)"
        )

    with col_pdf:
        # Geração PDF
        try:
            pdf_buffer = generate_pdf(apresentacao_markdown)
            st.download_button(
                label="Baixar PDF",
                data=pdf_buffer,
                file_name="SINAPSE_BR_IA_Apresentacao.pdf",
                mime="application/pdf",
                help="Baixa a página principal no formato PDF"
            )
        except Exception as e:
             st.error(f"Erro ao gerar PDF: {e}")
    # -----------------------------------------------------------------------------
    
# ---------------------------------
# Rodapé (permanece inalterado)
# ---------------------------------
st.markdown("---")
st.caption(
    f"Root detectado: `{PROJECT_ROOT}` • Imagens: `{IMG_DIR}` • "
    "Caminhos relativos compatíveis com execução local e Streamlit Cloud."
)
